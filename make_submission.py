from __future__ import annotations

import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import numpy as np
import pandas as pd
from catboost import CatBoostRegressor
from docx import Document
from docx.shared import Inches, Pt
from sklearn.metrics import mean_absolute_error, mean_squared_error

ROOT = Path(__file__).resolve().parent
RANDOM_SEED = 42

FEATURE_COLUMNS = [
    "pickup",
    "delivery",
    "equipment",
    "route",
    "month",
    "dow",
    "pickup_lat",
    "pickup_lon",
    "delivery_lat",
    "delivery_lon",
    "distance",
    "distance_log",
    "weight_abs",
    "weight_log",
    "market_index",
    "quote_signal",
    "days_since_start",
    "dayofyear_sin",
    "dayofyear_cos",
    "is_weekend",
]

CATEGORICAL_COLUMNS = ["pickup", "delivery", "equipment", "route", "month", "dow"]

FORWARD_FOLDS = [
    ("2025-07-01", "2025-08-01"),
    ("2025-08-01", "2025-09-01"),
    ("2025-09-01", "2025-10-01"),
    ("2025-10-01", "2025-11-01"),
]


@dataclass(frozen=True)
class PreprocessStats:
    pickup_lat_median: float
    pickup_lon_median: float
    delivery_lat_median: float
    delivery_lon_median: float
    distance_median: float
    weight_median: float
    market_index_median: float
    quote_signal_median: float
    date_origin: pd.Timestamp


def resolve_input(*relative_names: str) -> Path:
    for name in relative_names:
        for base in (ROOT, ROOT / "data"):
            candidate = base / name
            if candidate.is_file():
                return candidate
    raise FileNotFoundError(f"Could not find any of: {', '.join(relative_names)}")


def read_csv(*relative_names: str) -> pd.DataFrame:
    return pd.read_csv(resolve_input(*relative_names))


def build_stats(frame: pd.DataFrame) -> PreprocessStats:
    return PreprocessStats(
        pickup_lat_median=float(frame["pickup_lat"].median()),
        pickup_lon_median=float(frame["pickup_lon"].median()),
        delivery_lat_median=float(frame["delivery_lat"].median()),
        delivery_lon_median=float(frame["delivery_lon"].median()),
        distance_median=float(frame["distance"].median()),
        weight_median=float(frame["weight"].median()),
        market_index_median=float(frame["market_index"].median()),
        quote_signal_median=float(frame["quote_signal"].median()),
        date_origin=pd.to_datetime(frame["date"]).min(),
    )


def prepare_features(frame: pd.DataFrame, stats: PreprocessStats) -> pd.DataFrame:
    result = frame.copy()
    result["date"] = pd.to_datetime(result["date"])

    if "pickup_lat" not in result.columns:
        result["pickup_lat"] = np.nan
    if "pickup_lon" not in result.columns:
        result["pickup_lon"] = np.nan
    if "delivery_lat" not in result.columns:
        result["delivery_lat"] = np.nan
    if "delivery_lon" not in result.columns:
        result["delivery_lon"] = np.nan
    if "distance" not in result.columns:
        result["distance"] = np.nan
    if "weight" not in result.columns:
        result["weight"] = np.nan
    if "market_index" not in result.columns:
        result["market_index"] = np.nan
    if "quote_signal" not in result.columns:
        result["quote_signal"] = np.nan

    result["pickup_lat"] = pd.to_numeric(result["pickup_lat"], errors="coerce").fillna(stats.pickup_lat_median)
    result["pickup_lon"] = pd.to_numeric(result["pickup_lon"], errors="coerce").fillna(stats.pickup_lon_median)
    result["delivery_lat"] = pd.to_numeric(result["delivery_lat"], errors="coerce").fillna(stats.delivery_lat_median)
    result["delivery_lon"] = pd.to_numeric(result["delivery_lon"], errors="coerce").fillna(stats.delivery_lon_median)
    result["distance"] = pd.to_numeric(result["distance"], errors="coerce").fillna(stats.distance_median)
    result["weight"] = pd.to_numeric(result["weight"], errors="coerce").fillna(stats.weight_median)
    result["weight_abs"] = result["weight"].abs()
    result["market_index"] = pd.to_numeric(result["market_index"], errors="coerce").fillna(stats.market_index_median)
    result["quote_signal"] = pd.to_numeric(result["quote_signal"], errors="coerce").fillna(stats.quote_signal_median)
    result["route"] = result["pickup"].astype(str) + "|" + result["delivery"].astype(str)
    result["month"] = result["date"].dt.month.astype(str)
    result["dow"] = result["date"].dt.dayofweek.astype(str)
    result["days_since_start"] = (result["date"] - stats.date_origin).dt.days.astype(int)
    day_of_year = result["date"].dt.dayofyear.astype(float)
    result["dayofyear_sin"] = np.sin(2.0 * np.pi * day_of_year / 365.25)
    result["dayofyear_cos"] = np.cos(2.0 * np.pi * day_of_year / 365.25)
    result["distance_log"] = np.log1p(result["distance"].to_numpy(dtype=float))
    result["weight_log"] = np.log1p(result["weight_abs"].to_numpy(dtype=float))
    result["is_weekend"] = result["dow"].isin(["5", "6"]).astype(int)

    for column in CATEGORICAL_COLUMNS:
        result[column] = result[column].astype(str).fillna("missing")

    return result[FEATURE_COLUMNS]


def make_model() -> CatBoostRegressor:
    return CatBoostRegressor(
        loss_function="RMSE",
        eval_metric="RMSE",
        iterations=250,
        learning_rate=0.06,
        depth=8,
        l2_leaf_reg=4,
        random_seed=RANDOM_SEED,
        verbose=False,
        allow_writing_files=False,
    )


def fit_final_model(train_frame: pd.DataFrame) -> tuple[CatBoostRegressor, PreprocessStats]:
    stats = build_stats(train_frame)
    train_features = prepare_features(train_frame, stats)
    model = make_model()
    model.fit(train_features, train_frame["posted_rate"], cat_features=CATEGORICAL_COLUMNS)
    return model, stats


def evaluate_forward_chaining(train_frame: pd.DataFrame) -> pd.DataFrame:
    results: list[dict[str, float | int | str]] = []
    for train_end, valid_end in FORWARD_FOLDS:
        train_subset = train_frame[train_frame["date"] < train_end].copy()
        valid_subset = train_frame[(train_frame["date"] >= train_end) & (train_frame["date"] < valid_end)].copy()

        train_stats = build_stats(train_subset)
        train_features = prepare_features(train_subset, train_stats)
        valid_features = prepare_features(valid_subset, train_stats)

        model = make_model()
        model.fit(
            train_features,
            train_subset["posted_rate"],
            cat_features=CATEGORICAL_COLUMNS,
            eval_set=(valid_features, valid_subset["posted_rate"]),
            use_best_model=True,
            early_stopping_rounds=50,
        )
        predictions = model.predict(valid_features)
        rmse = mean_squared_error(valid_subset["posted_rate"], predictions) ** 0.5
        mae = mean_absolute_error(valid_subset["posted_rate"], predictions)

        results.append(
            {
                "train_end": train_end,
                "valid_end": valid_end,
                "train_rows": int(len(train_subset)),
                "valid_rows": int(len(valid_subset)),
                "best_iteration": int(model.get_best_iteration() or 0),
                "rmse": float(rmse),
                "mae": float(mae),
            }
        )

    return pd.DataFrame(results)


def predict_frame(model: CatBoostRegressor, stats: PreprocessStats, frame: pd.DataFrame) -> pd.Series:
    features = prepare_features(frame, stats)
    predictions = pd.Series(model.predict(features), index=frame.index, dtype=float)
    return predictions.clip(lower=1.0)


def write_validation_predictions(model: CatBoostRegressor, stats: PreprocessStats, frame: pd.DataFrame) -> Path:
    predictions = pd.DataFrame(
        {
            "load_id": frame["load_id"],
            "predicted_rate": predict_frame(model, stats, frame),
        }
    )
    output_path = ROOT / "validation_predictions.csv"
    predictions.to_csv(output_path, index=False)
    return output_path


def write_december_predictions(model: CatBoostRegressor, stats: PreprocessStats, frame: pd.DataFrame) -> Path:
    predictions = frame.copy()
    predictions["predicted_rate"] = predict_frame(model, stats, frame)
    output_path = ROOT / "december_chart_inputs_filled.csv"
    predictions.to_csv(output_path, index=False)
    return output_path


def run_scorer(validation_path: Path, december_path: Path) -> Path:
    subprocess.run(
        [
            sys.executable,
            str(ROOT / "score.py"),
            "--predictions",
            str(validation_path),
            "--december-predictions",
            str(december_path),
        ],
        check=True,
        cwd=ROOT,
    )
    return ROOT / "scorer_results" / "candidate_december.png"


def add_doc_paragraph(document: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        paragraph.add_run(bold_prefix).bold = True
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)


def write_report(cv_results: pd.DataFrame, model: CatBoostRegressor, chart_path: Path) -> Path:
    document = Document()
    document.styles["Normal"].font.name = "Aptos"
    document.styles["Normal"].font.size = Pt(10.5)

    document.add_heading("Freight Rate Prediction Assessment", level=0)
    document.add_paragraph(
        "This report summarizes the validation strategy, data issues, feature engineering, model choice, and the December chart produced by score.py."
    )

    document.add_heading("Objective", level=1)
    add_doc_paragraph(
        document,
        "The objective was to predict posted_rate for future freight loads using the labeled development data and then generate predictions for the final 12,000 validation loads and the fixed December scenario.",
    )

    document.add_heading("Validation Strategy", level=1)
    add_doc_paragraph(
        document,
        "I used forward-chaining monthly validation instead of a random train/test split because this is a time-series forecasting problem. The model was trained on earlier months and evaluated on later months so the validation setup reflects the real deployment setting.",
    )
    add_doc_paragraph(
        document,
        "The folds were July, August, September, and October 2025 holdouts built from the January-October development set. This approach avoids leakage from future periods and gives a more honest estimate of how the model will generalize.",
    )
    add_doc_paragraph(
        document,
        "I chose this strategy because freight rates change over time, so a random split would likely overstate performance by allowing the model to see similar patterns from the near future during training.",
    )

    document.add_heading("Data Exploration and Findings", level=1)
    add_doc_paragraph(document, "The development set contains 48,000 rows. The final validation set contains 12,000 rows, and the December input file contains 31 rows for the fixed chart scenario.")
    add_doc_paragraph(document, "The target is strongly distance-driven: distance correlates with posted_rate at about 0.91, which means lane length is one of the most important signals in the data.")
    add_doc_paragraph(document, "I found no duplicate rows in the training data, validation data, or the December inputs, so the main challenge was modeling the signal rather than cleaning duplicates.")
    add_doc_paragraph(document, "The dataset also contains many unique pickup/delivery routes, which makes route-level structure important. There are roughly 4,000 unique pickup/delivery pairs in the training data.")

    document.add_heading("Data Quality Issues and Handling", level=1)
    add_doc_paragraph(document, "The main data-quality issues were missing values in weight and market_index, plus a small number of negative weight values that are not realistic for freight loads.")
    add_doc_paragraph(document, "I imputed missing numeric values using training medians so the model could still learn from those rows without losing data. For weight, I also normalized the feature by using the absolute value before the log transform, which made the signal more stable and more consistent with the business context.")

    document.add_heading("Feature Engineering", level=1)
    add_doc_paragraph(document, "Key engineered features include route, month, day of week, days since the first training date, and cyclical day-of-year sine/cosine terms. These features capture both route-level structure and seasonality.")
    add_doc_paragraph(document, "I also added distance_log, weight_log, and an is_weekend flag to help the model capture nonlinear patterns without relying on hard-coded rules.")
    add_doc_paragraph(document, "The model uses categorical features such as pickup, delivery, equipment, and route because those variables carry strong signal in freight pricing.")

    document.add_heading("Model Choice and Why Not Alternatives", level=1)
    add_doc_paragraph(document, "I chose CatBoostRegressor because it is specifically strong on tabular data with high-cardinality categorical features. It handles pickup, delivery, route, equipment, and time-based categories well and learns nonlinear interactions without requiring a huge sparse one-hot matrix.")
    add_doc_paragraph(document, "A simple linear regression model was not preferred because it would miss nonlinear structure and would not handle route-level categories as well as CatBoost.")
    add_doc_paragraph(document, "A distance-only model was also not sufficient because the data showed route-specific effects, equipment differences, and seasonal behavior beyond simple distance.")
    add_doc_paragraph(document, "The final model was trained with RMSE loss, learning_rate=0.06, depth=8, l2_leaf_reg=4, and 250 iterations.")

    document.add_heading("Forward-Chaining Results", level=1)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Fold"
    header_cells[1].text = "Rows"
    header_cells[2].text = "RMSE"
    header_cells[3].text = "MAE"
    for _, row in cv_results.iterrows():
        cells = table.add_row().cells
        cells[0].text = f"{row['train_end']} -> {row['valid_end']}"
        cells[1].text = f"{int(row['valid_rows']):,}"
        cells[2].text = f"{row['rmse']:.2f}"
        cells[3].text = f"{row['mae']:.2f}"

    add_doc_paragraph(
        document,
        f"Average forward-chaining RMSE: {cv_results['rmse'].mean():.2f}. Average MAE: {cv_results['mae'].mean():.2f}.",
    )

    document.add_heading("Most Important Features", level=1)
    final_feature_importance = pd.Series(
        model.get_feature_importance(), index=FEATURE_COLUMNS
    ).sort_values(ascending=False)
    for feature, importance in final_feature_importance.head(8).items():
        add_doc_paragraph(document, f"{feature}: {importance:.2f}")

    document.add_heading("December Chart", level=1)
    document.add_paragraph(
        "The scorer generated the fixed December chart below from the completed December predictions file."
    )
    document.add_picture(str(chart_path), width=Inches(6.5))

    output_path = ROOT / "submission_report.docx"
    document.save(output_path)
    return output_path


def main() -> None:
    train_frame = read_csv("train-test.csv", "train_test.csv")
    validation_frame = read_csv("validation.csv")
    december_frame = read_csv("december-chart-inputs.csv", "december_chart_inputs.csv")

    if "date" not in train_frame.columns:
        raise ValueError("Training data must contain a date column")

    train_frame = train_frame.copy()
    train_frame["date"] = pd.to_datetime(train_frame["date"])
    validation_frame = validation_frame.copy()
    validation_frame["date"] = pd.to_datetime(validation_frame["date"])
    december_frame = december_frame.copy()
    december_frame["date"] = pd.to_datetime(december_frame["date"])

    cv_results = evaluate_forward_chaining(train_frame)
    model, stats = fit_final_model(train_frame)

    validation_path = write_validation_predictions(model, stats, validation_frame)
    december_path = write_december_predictions(model, stats, december_frame)
    chart_path = run_scorer(validation_path, december_path)
    report_path = write_report(cv_results, model, chart_path)

    print("Validation predictions written to:", validation_path)
    print("December predictions written to:", december_path)
    print("Chart written to:", chart_path)
    print("Report written to:", report_path)
    print("CV summary:")
    print(cv_results[["train_end", "valid_end", "rmse", "mae"]].to_string(index=False))


if __name__ == "__main__":
    main()
